import io
import json
import os
import sqlite3
import subprocess
import uuid
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone

import requests
from docx import Document
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

# WeasyPrint тянет системные библиотеки (Pango/Cairo). Если их нет на хосте — импорт падает
# с OSError. Раньше это падало на старте всего приложения и валило ВЕСЬ бэкенд, включая
# транскрипцию файлов. Теперь импортируем лениво, только когда реально нужен PDF-экспорт —
# остальной функционал продолжает работать, даже если PDF временно недоступен.
_weasyprint_error = None
try:
    from weasyprint import HTML
except Exception as e:  # noqa: BLE001
    HTML = None
    _weasyprint_error = str(e)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 3 * 1024 * 1024 * 1024  # 3 ГБ — с запасом под видеофайлы
# MVP: разрешаем любые origin (расширение работает без бэкенд-аутентификации).
# После публикации можно сузить до конкретного chrome-extension://<ID>.
CORS(app, resources={r"/api/*": {"origins": "*"}})

# Временные файлы загрузок (сырое видео/аудио до отправки в AssemblyAI) — НЕ в /data
# (тот volume только для sqlite и должен оставаться маленьким), а в эфемерном /tmp.
UPLOAD_TMP_DIR = os.environ.get("UPLOAD_TMP_DIR", "/tmp/uploads")
os.makedirs(UPLOAD_TMP_DIR, exist_ok=True)

# Сколько файлов обрабатываем параллельно в фоне (ffmpeg + заливка в AssemblyAI).
# Не привязано к числу воркеров gunicorn — это ограничивает именно CPU/сеть под капотом,
# остальные заявки просто встают в очередь пула, а не роняют сервер. Подбирай под реальные
# ресурсы Railway (больше vCPU/RAM — можно поднимать).
JOB_WORKERS = int(os.environ.get("JOB_WORKERS", "4"))
job_executor = ThreadPoolExecutor(max_workers=JOB_WORKERS, thread_name_prefix="job")

# Расширения, для которых имеет смысл сначала вырезать звук через ffmpeg — в видео
# звуковая дорожка обычно занимает не больше 10% веса файла, остальное тратить на
# заливку в AssemblyAI бессмысленно.
VIDEO_EXTENSIONS = {".mp4", ".mov", ".mkv", ".avi", ".webm", ".m4v", ".wmv", ".flv"}

ASSEMBLYAI_API_KEY = os.environ.get("ASSEMBLYAI_API_KEY")
ASSEMBLYAI_BASE = "https://api.assemblyai.com/v2"
SPEECH_MODEL = "universal-2"  # самая дешёвая модель AssemblyAI ($0.15/час) — фиксируем явно,
                            # чтобы себестоимость была предсказуемой, а не "что дадут по умолчанию"

FREE_LIMIT_SECONDS = 30 * 60  # 30 минут — держим то же значение, что и в extension/config.js

DODO_API_KEY = os.environ.get("DODO_PAYMENTS_API_KEY")
DODO_ENVIRONMENT = os.environ.get("DODO_ENVIRONMENT", "live_mode")  # "test_mode" на время тестов
DODO_WEBHOOK_KEY = os.environ.get("DODO_WEBHOOK_KEY")

# Три продукта-подписки, которые нужно создать в дашборде Dodo — id проставь в Railway Variables.
# minutes — это то, что должно быть настроено как "credits issued" на самом продукте в Dodo
# (Credit entitlement "minutes", precision 0) — тут только для отображения в /api/usage и UI.
PLANS = {
    "weekly": {"product_id": os.environ.get("DODO_PRODUCT_WEEKLY"), "label": "Weekly", "minutes": 600},
    "monthly": {"product_id": os.environ.get("DODO_PRODUCT_MONTHLY"), "label": "Monthly", "minutes": 1800},
    # Годовой биллится одним платежом в год — credits логично выдавать сразу на весь год
    # (1800/мес × 12), а не рассчитывать на помесячный reissue при годовом billing cycle.
    "yearly": {"product_id": os.environ.get("DODO_PRODUCT_YEARLY"), "label": "Yearly", "minutes": 1800 * 12},
}
# Событие, которое шлём в Meter после каждой завершённой транскрипции — должно ТОЧНО совпадать
# с Event Name, указанным при создании Meter в дашборде Dodo.
USAGE_EVENT_NAME = "transcription_completed"

_dodo_client = None
def dodo_client():
    global _dodo_client
    if _dodo_client is None:
        from dodopayments import DodoPayments
        _dodo_client = DodoPayments(bearer_token=DODO_API_KEY, environment=DODO_ENVIRONMENT)
    return _dodo_client

# На Railway подключи Volume и примонтируй его в /data, чтобы SQLite не терялась при redeploy.
DB_PATH = os.environ.get("DB_PATH", "/data/app.db")


def get_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = sqlite3.connect(DB_PATH, timeout=10)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")  # снижает блокировки при параллельных воркерах gunicorn
    return conn


def init_db():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS usage (
            device_id TEXT PRIMARY KEY,
            seconds_used REAL NOT NULL DEFAULT 0,
            updated_at TEXT
        )
    """)

    conn.execute("""
        CREATE TABLE IF NOT EXISTS customers (
            device_id TEXT PRIMARY KEY,
            dodo_customer_id TEXT NOT NULL,
            email TEXT,
            updated_at TEXT
        )
    """)

    conn.execute("""
        CREATE TABLE IF NOT EXISTS subscriptions (
            device_id TEXT PRIMARY KEY,
            subscription_id TEXT,
            plan TEXT,
            status TEXT,
            remaining_balance REAL,
            next_billing_date TEXT,
            email TEXT,
            updated_at TEXT
        )
    """)

    # Миграция: у более ранних версий таблиц customers/subscriptions не было колонки email —
    # без неё невозможно узнавать PRO-статус на новом устройстве по одному и тому же аккаунту.
    customers_cols = [row[1] for row in conn.execute("PRAGMA table_info(customers)").fetchall()]
    if "email" not in customers_cols:
        conn.execute("ALTER TABLE customers ADD COLUMN email TEXT")

    subs_cols = [row[1] for row in conn.execute("PRAGMA table_info(subscriptions)").fetchall()]
    if "email" not in subs_cols:
        conn.execute("ALTER TABLE subscriptions ADD COLUMN email TEXT")

    conn.execute("CREATE INDEX IF NOT EXISTS idx_customers_email ON customers(email)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_subscriptions_email ON subscriptions(email)")

    conn.execute("""
        CREATE TABLE IF NOT EXISTS jobs (
            job_id TEXT PRIMARY KEY,
            transcript_id TEXT,
            device_id TEXT NOT NULL,
            title TEXT,
            counted INTEGER NOT NULL DEFAULT 0,
            status TEXT NOT NULL DEFAULT 'received',
            error_detail TEXT,
            created_at TEXT,
            text TEXT,
            utterances TEXT,
            duration REAL
        )
    """)

    # Миграция со времён, когда была YouTube-ссылка: колонка source_type была NOT NULL,
    # без значения по умолчанию — INSERT в неё сейчас падает. CREATE TABLE IF NOT EXISTS
    # не трогает уже существующую таблицу, поэтому чиним руками, если колонка ещё жива.
    existing_cols = [row[1] for row in conn.execute("PRAGMA table_info(jobs)").fetchall()]
    if "source_type" in existing_cols:
        conn.execute("ALTER TABLE jobs RENAME TO jobs_old")
        conn.execute("""
            CREATE TABLE jobs (
                job_id TEXT PRIMARY KEY,
                transcript_id TEXT,
                device_id TEXT NOT NULL,
                title TEXT,
                counted INTEGER NOT NULL DEFAULT 0,
                status TEXT NOT NULL DEFAULT 'received',
                error_detail TEXT,
                created_at TEXT,
                text TEXT,
                utterances TEXT,
                duration REAL
            )
        """)
        old_cols = [row[1] for row in conn.execute("PRAGMA table_info(jobs_old)").fetchall()]
        # У старых строк job_id ещё не было — раньше единственным идентификатором был
        # transcript_id (его же фронтенд уже хранит как "id" в локальной истории), поэтому
        # переиспользуем его как job_id, чтобы старые ссылки на экспорт не сломались.
        target_cols = ["job_id", "transcript_id", "device_id", "title", "counted", "status", "created_at", "text", "utterances", "duration"]
        select_list = ", ".join([
            "transcript_id", "transcript_id",
            *(c if c in old_cols else "NULL" for c in ["device_id", "title", "counted"]),
            "'processing'",
            *(c if c in old_cols else "NULL" for c in ["created_at", "text", "utterances", "duration"]),
        ])
        conn.execute(f"INSERT INTO jobs ({', '.join(target_cols)}) SELECT {select_list} FROM jobs_old")
        conn.execute("DROP TABLE jobs_old")

    # Отдельная миграция: у более ранних версий таблицы jobs (уже без source_type) primary
    # key был transcript_id, а не job_id — тот же перенос данных, что и выше.
    jobs_cols = [row[1] for row in conn.execute("PRAGMA table_info(jobs)").fetchall()]
    if "job_id" not in jobs_cols:
        conn.execute("ALTER TABLE jobs RENAME TO jobs_old")
        conn.execute("""
            CREATE TABLE jobs (
                job_id TEXT PRIMARY KEY,
                transcript_id TEXT,
                device_id TEXT NOT NULL,
                title TEXT,
                counted INTEGER NOT NULL DEFAULT 0,
                status TEXT NOT NULL DEFAULT 'received',
                error_detail TEXT,
                created_at TEXT,
                text TEXT,
                utterances TEXT,
                duration REAL
            )
        """)
        conn.execute("""
            INSERT INTO jobs (job_id, transcript_id, device_id, title, counted, status, created_at, text, utterances, duration)
            SELECT transcript_id, transcript_id, device_id, title, counted, 'processing', created_at, text, utterances, duration
            FROM jobs_old
        """)
        conn.execute("DROP TABLE jobs_old")

    conn.execute("CREATE INDEX IF NOT EXISTS idx_jobs_transcript_id ON jobs(transcript_id)")


    # Более старая миграция: колонка paid_seconds_balance из версии с разовыми платежами ($1/час)
    # больше не нужна — теперь баланс живёт в Dodo и приходит через вебхук подписки.
    usage_cols = [row[1] for row in conn.execute("PRAGMA table_info(usage)").fetchall()]
    if "paid_seconds_balance" in usage_cols:
        conn.execute("ALTER TABLE usage RENAME TO usage_old")
        conn.execute("""
            CREATE TABLE usage (
                device_id TEXT PRIMARY KEY,
                seconds_used REAL NOT NULL DEFAULT 0,
                updated_at TEXT
            )
        """)
        conn.execute("INSERT INTO usage (device_id, seconds_used, updated_at) SELECT device_id, seconds_used, updated_at FROM usage_old")
        conn.execute("DROP TABLE usage_old")

    conn.commit()
    conn.close()


init_db()


def get_usage_row(device_id):
    conn = get_db()
    row = conn.execute("SELECT seconds_used FROM usage WHERE device_id=?", (device_id,)).fetchone()
    conn.close()
    return {"seconds_used": row["seconds_used"] if row else 0}


def get_subscription(device_id, email=None):
    """PRO-статус привязан к почте: сначала смотрим подписку этого устройства, а если её
    нет (или она неактивна) и передан email — ищем активную подписку того же аккаунта на
    ЛЮБОМ другом устройстве. Найдя — сразу копируем её на текущий device_id, чтобы новое
    устройство/переустановка сразу увидели PRO без ожидания вебхука."""
    conn = get_db()
    row = conn.execute("SELECT * FROM subscriptions WHERE device_id=?", (device_id,)).fetchone()
    conn.close()

    if row and row["status"] == "active":
        return row

    if email:
        conn = get_db()
        email_row = conn.execute(
            "SELECT * FROM subscriptions WHERE email=? AND status='active' ORDER BY updated_at DESC LIMIT 1",
            (email,),
        ).fetchone()
        conn.close()
        if email_row and email_row["device_id"] != device_id:
            upsert_subscription(
                device_id,
                email_row["subscription_id"],
                email_row["plan"],
                email_row["status"],
                email_row["remaining_balance"],
                email_row["next_billing_date"],
                email,
            )
            dodo_customer_id = get_dodo_customer_id(email_row["device_id"])
            if dodo_customer_id:
                upsert_customer(device_id, dodo_customer_id, email)
            conn = get_db()
            row = conn.execute("SELECT * FROM subscriptions WHERE device_id=?", (device_id,)).fetchone()
            conn.close()
            return row

    return row


def has_credit(device_id, email=None):
    """Бесплатные 30 минут (на устройство) ИЛИ активная подписка — своя или найденная по
    почте на другом устройстве (превышение лимита Dodo биллит сама — поэтому наличие активной
    подписки достаточно, отдельно считать остаток не нужно)."""
    if get_usage_row(device_id)["seconds_used"] < FREE_LIMIT_SECONDS:
        return True
    sub = get_subscription(device_id, email)
    return bool(sub and sub["status"] == "active")


def add_free_seconds_used(device_id, seconds):
    conn = get_db()
    conn.execute("""
        INSERT INTO usage (device_id, seconds_used, updated_at)
        VALUES (?, ?, ?)
        ON CONFLICT(device_id) DO UPDATE SET
            seconds_used = seconds_used + excluded.seconds_used,
            updated_at = excluded.updated_at
    """, (device_id, seconds, datetime.now(timezone.utc).isoformat()))
    conn.commit()
    conn.close()


def upsert_customer(device_id, dodo_customer_id, email=None):
    conn = get_db()
    conn.execute("""
        INSERT INTO customers (device_id, dodo_customer_id, email, updated_at)
        VALUES (?, ?, ?, ?)
        ON CONFLICT(device_id) DO UPDATE SET
            dodo_customer_id = excluded.dodo_customer_id,
            email = COALESCE(excluded.email, customers.email),
            updated_at = excluded.updated_at
    """, (device_id, dodo_customer_id, email, datetime.now(timezone.utc).isoformat()))
    conn.commit()
    conn.close()


def get_dodo_customer_id(device_id, email=None):
    conn = get_db()
    row = conn.execute("SELECT dodo_customer_id FROM customers WHERE device_id=?", (device_id,)).fetchone()
    if not row and email:
        email_row = conn.execute(
            "SELECT dodo_customer_id FROM customers WHERE email=? ORDER BY updated_at DESC LIMIT 1",
            (email,),
        ).fetchone()
        if email_row:
            conn.close()
            upsert_customer(device_id, email_row["dodo_customer_id"], email)
            return email_row["dodo_customer_id"]
    conn.close()
    return row["dodo_customer_id"] if row else None


def upsert_subscription(device_id, subscription_id, plan, status, remaining_balance, next_billing_date, email=None):
    conn = get_db()
    conn.execute("""
        INSERT INTO subscriptions (device_id, subscription_id, plan, status, remaining_balance, next_billing_date, email, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(device_id) DO UPDATE SET
            subscription_id = excluded.subscription_id,
            plan = excluded.plan,
            status = excluded.status,
            remaining_balance = excluded.remaining_balance,
            next_billing_date = excluded.next_billing_date,
            email = COALESCE(excluded.email, subscriptions.email),
            updated_at = excluded.updated_at
    """, (device_id, subscription_id, plan, status, remaining_balance, next_billing_date, email, datetime.now(timezone.utc).isoformat()))
    conn.commit()
    conn.close()


def send_usage_event(device_id, transcript_id, duration_seconds):
    """Шлём событие в Dodo, чтобы Meter списал минуты с баланса подписки (или выставил
    overage, если баланс исчерпан). Если у устройства ещё нет клиента Dodo (никогда не
    оформляло подписку — всё ещё на бесплатных 30 минутах), просто ничего не отправляем."""
    dodo_customer_id = get_dodo_customer_id(device_id)
    if not dodo_customer_id:
        return
    minutes = max(1, int((duration_seconds + 59) // 60))  # округляем вверх до целой минуты
    try:
        dodo_client().usage_events.ingest(events=[{
            "event_id": transcript_id,  # тот же id при retry — по идее не задвоит событие
            "customer_id": dodo_customer_id,
            "event_name": USAGE_EVENT_NAME,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "metadata": {"minutes": str(minutes)},
        }])
    except Exception as e:  # noqa: BLE001
        print(f"[warn] failed to send usage event for {transcript_id}: {e}")


def save_job(job_id, device_id, title):
    conn = get_db()
    conn.execute(
        "INSERT INTO jobs (job_id, device_id, title, status, created_at) VALUES (?, ?, ?, ?, ?)",
        (job_id, device_id, title, "received", datetime.now(timezone.utc).isoformat()),
    )
    conn.commit()
    conn.close()


def get_job(job_id):
    conn = get_db()
    row = conn.execute("SELECT * FROM jobs WHERE job_id=?", (job_id,)).fetchone()
    conn.close()
    return row


def update_job_status(job_id, status, error_detail=None):
    conn = get_db()
    conn.execute("UPDATE jobs SET status=?, error_detail=? WHERE job_id=?", (status, error_detail, job_id))
    conn.commit()
    conn.close()


def set_job_transcript_id(job_id, transcript_id):
    conn = get_db()
    conn.execute("UPDATE jobs SET transcript_id=?, status=? WHERE job_id=?", (transcript_id, "processing", job_id))
    conn.commit()
    conn.close()


def mark_job_counted(job_id):
    conn = get_db()
    conn.execute("UPDATE jobs SET counted=1 WHERE job_id=?", (job_id,))
    conn.commit()
    conn.close()


def save_job_result(job_id, text, utterances, duration):
    conn = get_db()
    conn.execute(
        "UPDATE jobs SET text=?, utterances=?, duration=? WHERE job_id=?",
        (text, json.dumps(utterances, ensure_ascii=False), duration, job_id),
    )
    conn.commit()
    conn.close()


def is_video_file(filename):
    ext = os.path.splitext(filename or "")[1].lower()
    return ext in VIDEO_EXTENSIONS


def extract_audio(input_path, output_path):
    """Вырезаем только звук из видео через ffmpeg — экономит и память (обрабатывается
    потоково, не грузится в RAM целиком), и трафик/время до AssemblyAI, т.к. видео-дорожка
    (обычно основной вес файла) для транскрипции не нужна."""
    cmd = [
        "ffmpeg", "-y", "-i", input_path,
        "-vn",                  # без видео
        "-ac", "1",              # моно — распознаванию речи стерео не нужно
        "-ar", "16000",          # AssemblyAI всё равно ресэмплит в 16kHz — сразу отдаём в этом виде
        "-c:a", "aac", "-b:a", "96k",
        output_path,
    ]
    result = subprocess.run(cmd, capture_output=True, timeout=1800)  # 30 минут максимум на конвертацию
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg failed: {result.stderr.decode(errors='ignore')[-500:]}")


def process_transcription_job(job_id, filepath, original_filename):
    """Работает в фоновом пуле потоков — не держит воркер gunicorn всё время загрузки/
    конвертации/заливки в AssemblyAI, поэтому другие пользователи не встают в очередь
    позади одной большой задачи."""
    extracted_path = None
    try:
        audio_path = filepath
        if is_video_file(original_filename):
            update_job_status(job_id, "extracting")
            extracted_path = f"{filepath}.audio.m4a"
            extract_audio(filepath, extracted_path)
            audio_path = extracted_path

        update_job_status(job_id, "uploading")
        with open(audio_path, "rb") as f:
            # requests стримит файловый объект по частям, а не грузит всё содержимое в
            # память разом — важно для аудио, вырезанного из многогигабайтного видео.
            upload_res = requests.post(
                f"{ASSEMBLYAI_BASE}/upload",
                headers=assemblyai_headers(),
                data=f,
            )
        upload_res.raise_for_status()
        audio_url = upload_res.json()["upload_url"]

        transcript_res = requests.post(
            f"{ASSEMBLYAI_BASE}/transcript",
            headers=assemblyai_headers(),
            json={
                "audio_url": audio_url,
                "speaker_labels": True,
                "speech_models": [SPEECH_MODEL],
                # Без этого AssemblyAI по умолчанию считает язык английским (en_us) и
                # "подгоняет" под него любую речь. language_detection сама определяет
                # язык (в т.ч. смену языка внутри одного файла — code switching).
                "language_detection": True,
            },
        )
        transcript_res.raise_for_status()
        transcript_id = transcript_res.json()["id"]
        set_job_transcript_id(job_id, transcript_id)
    except Exception as e:  # noqa: BLE001
        print(f"[error] job {job_id} failed: {e}")
        update_job_status(job_id, "error", error_detail=str(e)[:500])
    finally:
        for p in (filepath, extracted_path):
            if p and os.path.exists(p):
                try:
                    os.remove(p)
                except OSError:
                    pass


def _run_job_safely(job_id, filepath, original_filename):
    """Обёртка на случай непредвиденной ошибки ВНЕ try/except внутри самой задачи (как
    только что произошло с случайно затёртой build_diarized_text). Без этой страховки
    исключение в фоновом потоке ThreadPoolExecutor молча проглатывается (Future никто
    не читает), и задача зависает в статусе "uploading" навсегда — ни ошибки у
    пользователя, ни следа в логах."""
    try:
        process_transcription_job(job_id, filepath, original_filename)
    except Exception as e:  # noqa: BLE001
        print(f"[error] unhandled failure in job {job_id}: {e}")
        try:
            update_job_status(job_id, "error", error_detail=str(e)[:500])
        except Exception:  # noqa: BLE001
            pass


def build_diarized_text(data):
    """Собираем текст с метками спикеров, если AssemblyAI вернул диаризацию."""
    utterances = data.get("utterances") or []
    if not utterances:
        return data.get("text", ""), []
    lines = [f"Speaker {u['speaker']}: {u['text']}" for u in utterances]
    return "\n\n".join(lines), utterances


def ms_to_srt_timestamp(ms):
    total_seconds, millis = divmod(int(ms), 1000)
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d},{millis:03d}"


def build_srt(utterances, fallback_text, duration_seconds):
    if not utterances:
        # Нет данных о таймингах реплик — отдаём один блок на весь ролик.
        end_ts = ms_to_srt_timestamp((duration_seconds or 0) * 1000)
        return f"1\n00:00:00,000 --> {end_ts}\n{fallback_text}\n"
    blocks = []
    for i, u in enumerate(utterances, start=1):
        start_ts = ms_to_srt_timestamp(u["start"])
        end_ts = ms_to_srt_timestamp(u["end"])
        blocks.append(f"{i}\n{start_ts} --> {end_ts}\nSpeaker {u['speaker']}: {u['text']}\n")
    return "\n".join(blocks)


def build_docx(title, text, utterances):
    doc = Document()
    doc.add_heading(title, level=1)
    if utterances:
        for u in utterances:
            p = doc.add_paragraph()
            run = p.add_run(f"Speaker {u['speaker']}: ")
            run.bold = True
            p.add_run(u["text"])
    else:
        for para in text.split("\n"):
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf(title, text, utterances):
    if HTML is None:
        raise RuntimeError(f"PDF export unavailable: {_weasyprint_error}")
    if utterances:
        body = "".join(
            f"<p><b>Speaker {u['speaker']}:</b> {u['text']}</p>" for u in utterances
        )
    else:
        body = "".join(f"<p>{para}</p>" for para in text.split("\n") if para.strip())
    html = f"""
    <html><head><meta charset="utf-8"><style>
        body {{ font-family: sans-serif; font-size: 13px; line-height: 1.5; }}
        h1 {{ font-size: 18px; }}
    </style></head>
    <body><h1>{title}</h1>{body}</body></html>
    """
    buf = io.BytesIO()
    HTML(string=html).write_pdf(buf)
    buf.seek(0)
    return buf


def assemblyai_headers():
    return {"authorization": ASSEMBLYAI_API_KEY}


@app.route("/api/usage", methods=["GET"])
def usage():
    device_id = request.args.get("device_id", "")
    email = request.args.get("email") or None
    u = get_usage_row(device_id)
    sub = get_subscription(device_id, email)
    return jsonify({
        "seconds_used": u["seconds_used"],
        "limit_seconds": FREE_LIMIT_SECONDS,
        "subscription": {
            "plan": sub["plan"],
            "status": sub["status"],
            "remaining_balance": sub["remaining_balance"],
            "next_billing_date": sub["next_billing_date"],
        } if sub else None,
    })


@app.route("/api/transcribe/file", methods=["POST"])
def transcribe_file():
    device_id = request.form.get("device_id", "")
    email = request.form.get("email") or None
    if not has_credit(device_id, email):
        return jsonify({"error": "limit_reached"}), 402

    uploaded = request.files.get("file")
    if not uploaded:
        return jsonify({"error": "no_file"}), 400

    job_id = str(uuid.uuid4())
    safe_name = secure_filename(uploaded.filename or "audio") or "audio"
    filepath = os.path.join(UPLOAD_TMP_DIR, f"{job_id}_{safe_name}")
    # .save() пишет на диск потоково (чанками), а НЕ грузит весь файл в RAM разом —
    # критично для видео на 2-3 ГБ при большом числе одновременных пользователей.
    uploaded.save(filepath)

    save_job(job_id, device_id, uploaded.filename)

    # Вырезание звука (для видео) и заливка в AssemblyAI уходят в фоновый пул потоков —
    # запрос не держит воркер gunicorn всё это время, поэтому другие пользователи не
    # встают в очередь позади одной большой задачи.
    job_executor.submit(_run_job_safely, job_id, filepath, uploaded.filename)

    return jsonify({"job_id": job_id})


@app.route("/api/status/<job_id>", methods=["GET"])
def status(job_id):
    device_id = request.args.get("device_id", "")
    email = request.args.get("email") or None
    job = get_job(job_id)
    if not job:
        return jsonify({"error": "job_not_found"}), 404

    if job["status"] == "error":
        return jsonify({"status": "error", "detail": job["error_detail"]})

    if not job["transcript_id"]:
        # Ещё готовим файл на нашей стороне (received / extracting / uploading) —
        # транскрипта в AssemblyAI пока не существует, спрашивать там нечего.
        return jsonify({"status": job["status"]})

    transcript_id = job["transcript_id"]
    res = requests.get(f"{ASSEMBLYAI_BASE}/transcript/{transcript_id}", headers=assemblyai_headers())
    res.raise_for_status()
    data = res.json()

    aai_status = data["status"]  # queued | processing | completed | error

    if aai_status == "completed":
        duration = data.get("audio_duration") or 0
        diarized_text, utterances = build_diarized_text(data)
        save_job_result(job_id, diarized_text, utterances, duration)
        if not job["counted"]:
            sub = get_subscription(device_id, email)
            is_subscribed = bool(sub and sub["status"] == "active")

            if is_subscribed:
                # Активная подписка — ВСЯ длительность списывается с неё через Dodo.
                # Раньше здесь сначала расходовались бесплатные 30 минут (общий счётчик
                # seconds_used) и только остаток сверху уходил в send_usage_event — из-за
                # этого у только что оформивших PRO первые 30 минут вообще не уменьшали
                # баланс подписки, хотя человек уже платит. Для подписчиков бесплатный
                # бакет больше не участвует вовсе.
                if duration > 0:
                    send_usage_event(device_id, transcript_id, duration)
            else:
                # Без подписки — как раньше, списываем из бесплатных 30 минут
                # (has_credit уже не пропустил бы сюда, если бы они кончились).
                u = get_usage_row(device_id)
                remaining_free = max(0, FREE_LIMIT_SECONDS - u["seconds_used"])
                from_free = min(duration, remaining_free)
                if from_free > 0:
                    add_free_seconds_used(device_id, from_free)

            mark_job_counted(job_id)
        return jsonify({
            "status": "completed",
            "text": diarized_text,
            "audio_duration": duration,
            "title": job["title"],
        })

    if aai_status == "error":
        update_job_status(job_id, "error", error_detail=data.get("error"))
        return jsonify({"status": "error", "detail": data.get("error")})

    return jsonify({"status": aai_status})


@app.route("/api/export/<job_id>", methods=["GET"])
def export(job_id):
    fmt = request.args.get("format", "txt")
    job = get_job(job_id)
    if not job:
        return jsonify({"error": "job_not_found"}), 404

    text = job["text"] or ""
    title = (job["title"] or "transcript").rsplit(".", 1)[0]
    utterances = json.loads(job["utterances"]) if job["utterances"] else []
    safe_name = "".join(c for c in title if c.isalnum() or c in " _-").strip() or "transcript"

    if fmt == "txt":
        buf = io.BytesIO(text.encode("utf-8"))
        return send_file(buf, mimetype="text/plain", as_attachment=True, download_name=f"{safe_name}.txt")

    if fmt == "srt":
        srt_content = build_srt(utterances, text, job["duration"] or 0)
        buf = io.BytesIO(srt_content.encode("utf-8"))
        return send_file(buf, mimetype="application/x-subrip", as_attachment=True, download_name=f"{safe_name}.srt")

    if fmt == "docx":
        buf = build_docx(title, text, utterances)
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{safe_name}.docx",
        )

    if fmt == "pdf":
        try:
            buf = build_pdf(title, text, utterances)
        except RuntimeError as e:
            return jsonify({"error": "pdf_unavailable", "detail": str(e)}), 503
        return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=f"{safe_name}.pdf")

    return jsonify({"error": "unknown_format"}), 400


@app.route("/api/subscribe", methods=["POST"])
def subscribe():
    data = request.get_json(force=True)
    device_id = data.get("device_id", "")
    plan = data.get("plan", "")
    email = data.get("email")  # из Google-аккаунта, если пользователь уже вошёл
    name = data.get("name")

    if not device_id:
        return jsonify({"error": "no_device_id"}), 400
    # Вход через Google обязателен перед оплатой (делает это extension UI) — без email
    # PRO-статус невозможно привязать к аккаунту и восстановить на другом устройстве.
    if not email:
        return jsonify({"error": "email_required"}), 400
    if plan not in PLANS:
        return jsonify({"error": "invalid_plan"}), 400
    product_id = PLANS[plan]["product_id"]
    if not DODO_API_KEY or not product_id:
        return jsonify({"error": "payments_not_configured"}), 503

    checkout_kwargs = dict(
        product_cart=[{"product_id": product_id, "quantity": 1}],
        return_url=f"{request.url_root.rstrip('/')}/payment-success",
        metadata={"device_id": device_id, "plan": plan, "email": email},
        # Автозаполнение почты на странице оплаты Dodo — пользователь ничего не вводит руками.
        customer={"email": email, "name": name or email},
    )

    try:
        session = dodo_client().checkout_sessions.create(**checkout_kwargs)
    except Exception as e:  # noqa: BLE001
        return jsonify({"error": "dodo_error", "detail": str(e)}), 502

    checkout_url = session.checkout_url
    if not checkout_url:
        return jsonify({"error": "no_checkout_url"}), 502

    return jsonify({"checkout_url": checkout_url})


@app.route("/api/customer-portal", methods=["GET"])
def customer_portal():
    device_id = request.args.get("device_id", "")
    email = request.args.get("email") or None
    dodo_customer_id = get_dodo_customer_id(device_id, email)
    if not dodo_customer_id:
        return jsonify({"error": "no_customer"}), 404
    try:
        session = dodo_client().customers.customer_portal.create(dodo_customer_id)
    except Exception as e:  # noqa: BLE001
        return jsonify({"error": "dodo_error", "detail": str(e)}), 502
    return jsonify({"url": session.link})


def _extract_balance_after(data):
    """Достаём остаток из полезной нагрузки credit.* вебхука. Точное имя поля Dodo нигде
    явно не документирует построчно, поэтому перебираем правдоподобные варианты — так
    надёжнее, чем гадать один конкретный ключ и молча ничего не находить."""
    candidates = (data, data.get("credit") or {}, data.get("entitlement") or {}, data.get("credit_entitlement") or {})
    for obj in candidates:
        if not isinstance(obj, dict):
            continue
        for key in ("balance_after", "balance", "remaining_balance", "current_balance"):
            if obj.get(key) is not None:
                try:
                    return float(obj[key])
                except (TypeError, ValueError):
                    continue
    return None


def _extract_customer_id(data):
    customer = data.get("customer")
    if isinstance(customer, dict) and customer.get("customer_id"):
        return customer["customer_id"]
    return data.get("customer_id")


def update_balance_for_customer(dodo_customer_id, remaining_balance):
    """Списание могло случиться на устройстве, где подписка НЕ была оформлена (мы же
    привязываем PRO к email на любое устройство) — поэтому обновляем remaining_balance
    сразу у всех device_id этого customer_id, а не только у того, что в metadata."""
    conn = get_db()
    conn.execute(
        """UPDATE subscriptions SET remaining_balance=?, updated_at=?
           WHERE device_id IN (SELECT device_id FROM customers WHERE dodo_customer_id=?)""",
        (remaining_balance, datetime.now(timezone.utc).isoformat(), dodo_customer_id),
    )
    conn.commit()
    conn.close()


@app.route("/api/webhooks/dodo", methods=["POST"])
def dodo_webhook():
    raw_body = request.get_data()
    if DODO_WEBHOOK_KEY:
        try:
            from standardwebhooks.webhooks import Webhook
            wh = Webhook(DODO_WEBHOOK_KEY)
            wh.verify(raw_body, {
                "webhook-id": request.headers.get("webhook-id", ""),
                "webhook-signature": request.headers.get("webhook-signature", ""),
                "webhook-timestamp": request.headers.get("webhook-timestamp", ""),
            })
        except Exception as e:  # noqa: BLE001
            return jsonify({"error": "invalid_signature", "detail": str(e)}), 401

    payload = request.get_json(force=True)
    event_type = payload.get("type", "")

    # Обрабатываем ЛЮБОЕ событие subscription.* одинаково — в разных версиях доков Dodo
    # встречаются subscription.active / subscription.updated / subscription.renewed и т.д.,
    # но структура полезной нагрузки (data) одна и та же, включая remaining_balance.
    if event_type.startswith("subscription."):
        sub_data = payload.get("data", {})
        metadata = sub_data.get("metadata", {})
        device_id = metadata.get("device_id")
        plan = metadata.get("plan")
        customer = sub_data.get("customer", {})
        dodo_customer_id = customer.get("customer_id") if isinstance(customer, dict) else None
        # email из metadata (проставили сами в /api/subscribe) либо, если Dodo его вернул,
        # из объекта customer — нужен, чтобы PRO-статус узнавался на любом устройстве этого аккаунта.
        email = metadata.get("email") or (customer.get("email") if isinstance(customer, dict) else None)

        if device_id and dodo_customer_id:
            upsert_customer(device_id, dodo_customer_id, email)

            remaining_balance = None
            credit_cart = sub_data.get("credit_entitlement_cart") or []
            if credit_cart:
                try:
                    remaining_balance = float(credit_cart[0].get("remaining_balance"))
                except (TypeError, ValueError):
                    remaining_balance = None

            upsert_subscription(
                device_id,
                sub_data.get("subscription_id"),
                plan,
                sub_data.get("status"),
                remaining_balance,
                sub_data.get("next_billing_date"),
                email,
            )

    # credit.* — отдельная категория событий (не subscription.*!) — именно она приходит
    # при каждом списании через meter (наши send_usage_event) и при начислении новых
    # credits. Раньше мы её вообще не слушали, поэтому remaining_balance обновлялся
    # только в момент оформления/продления подписки и не менялся при использовании.
    elif event_type.startswith("credit."):
        data = payload.get("data", {})
        dodo_customer_id = _extract_customer_id(data)
        remaining_balance = _extract_balance_after(data)
        if dodo_customer_id and remaining_balance is not None:
            update_balance_for_customer(dodo_customer_id, remaining_balance)
        else:
            # Не смогли распарсить — печатаем сырой payload, чтобы разобрать точные имена
            # полей по реальному событию из логов, а не гадать по документации.
            print(f"[warn] credit webhook '{event_type}' — couldn't parse customer/balance: {json.dumps(data)[:500]}")

    return jsonify({"ok": True})


@app.route("/payment-success", methods=["GET"])
def payment_success():
    return """
    <html><body style="font-family:sans-serif;text-align:center;padding:60px 20px;">
        <h2>Оплата прошла успешно ✓</h2>
        <p>Можно закрыть эту вкладку и вернуться в расширение — баланс обновится за пару секунд.</p>
    </body></html>
    """


@app.route("/", methods=["GET"])
def health():
    return jsonify({"ok": True, "pdf_available": HTML is not None})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
